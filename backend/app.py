import os
import re
import logging
import tempfile
import pdfkit
import json
import bleach
import string
import random
import sqlite3
import requests
import stripe
import time
import socket
import ipaddress
import threading
import io
import uuid
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from concurrent.futures import ThreadPoolExecutor, as_completed

from flask import Flask, request, jsonify, send_file, redirect, current_app, render_template_string, abort
from flask_cors import CORS
from urllib.parse import urlparse
from dotenv import load_dotenv
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail, Email
from markdown2 import markdown
from pydub import AudioSegment
from docx import Document

# Optional Rate Limiting
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address

    USE_LIMITER = True
except ImportError:
    USE_LIMITER = False

# Load env vars
load_dotenv()

# Logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Required env vars
SENDGRID_API_KEY = os.getenv('SENDGRID_API_KEY')
RECEIVER_EMAIL = os.getenv('RECEIVER_EMAIL')
VERIFIED_SENDER = os.getenv('VERIFIED_SENDER')
DISCORD_WEBHOOK_URL = os.getenv('DISCORD_WEBHOOK_URL')
STRIPE_SECRET_KEY = os.getenv('STRIPE_SECRET_KEY')
STRIPE_WEBHOOK_SECRET = os.getenv('STRIPE_WEBHOOK_SECRET')
FRONTEND_BASE_URL = os.getenv("FRONTEND_BASE_URL", "http://localhost:5173")
EBAY_OAUTH_TOKEN = os.getenv('EBAY_OAUTH_TOKEN')
EBAY_CLIENT_ID = os.getenv('EBAY_CLIENT_ID')
EBAY_CLIENT_SECRET = os.getenv('EBAY_CLIENT_SECRET')
EBAY_REFRESH_TOKEN = os.getenv('EBAY_REFRESH_TOKEN')

if not SENDGRID_API_KEY:
    raise RuntimeError("SENDGRID_API_KEY is required")
if not RECEIVER_EMAIL:
    raise RuntimeError("RECEIVER_EMAIL is required")
if not VERIFIED_SENDER:
    raise RuntimeError("VERIFIED_SENDER is required")
if not STRIPE_SECRET_KEY:
    raise RuntimeError("STRIPE_SECRET_KEY is required")
if not STRIPE_WEBHOOK_SECRET:
    raise RuntimeError("STRIPE_WEBHOOK_SECRET is required")

# Flask App
app = Flask(__name__)
CORS(app)

# Rate Limiter
if USE_LIMITER:
    limiter = Limiter(get_remote_address, app=app, default_limits=["5 per minute"])

# Email Validation
EMAIL_REGEX = re.compile(r'^[^@\s]+@[^@\s]+\.[^@\s]+$')

MAX_FILE_SIZE_MB = 512
ALLOWED_MIME_TYPES = {
    'text/plain', 'text/markdown', 'text/html',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'audio/mpeg', 'audio/wav', 'audio/ogg'
}

with open("cards.json", "r", encoding="utf-8") as f:
    ALL_CARDS = json.load(f)["data"]
CARD_LOOKUP = {str(card["id"]): card for card in ALL_CARDS}

generation_jobs = {}
# Store token and its expiry in memory
ebay_access_token = None
ebay_token_expiry = 0


def is_valid_email(email): return EMAIL_REGEX.match(email) is not None


def update_cards_periodically():
    import time
    while True:
        try:
            logger.info("Refreshing cards.json from YGOPRODeck...")
            resp = requests.get("https://db.ygoprodeck.com/api/v7/cardinfo.php", timeout=15)
            if resp.status_code == 200:
                with open("cards.json", "w", encoding="utf-8") as f:
                    f.write(resp.text)
                global ALL_CARDS, CARD_LOOKUP
                ALL_CARDS = json.loads(resp.text)["data"]
                CARD_LOOKUP = {str(card["id"]): card for card in ALL_CARDS}
                logger.info("Updated cards.json successfully with %d cards.", len(ALL_CARDS))
            else:
                logger.warning("Failed to update cards.json (status code %s)", resp.status_code)
        except Exception as e:
            logger.exception("Error during card update")
        time.sleep(21600)  # Sleep 6 hours


# ----------------------------------------------------------------------------------------------------------------------
# Database Initialisation
# ----------------------------------------------------------------------------------------------------------------------
DATABASE = 'shortener.db'


def get_db():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with get_db() as db:
        db.execute('''
            CREATE TABLE IF NOT EXISTS short_urls (
                code TEXT PRIMARY KEY,
                url TEXT NOT NULL UNIQUE,
                created TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')


init_db()


# ----------------------------------------------------------------------------------------------------------------------
# Contact Form Endpoint
# ----------------------------------------------------------------------------------------------------------------------
@app.route('/api/contact', methods=['POST'])
@limiter.limit('5 per minute') if USE_LIMITER else (lambda f: f)
def contact():
    logger.info("POST /api/contact")
    data = request.get_json(force=True, silent=True)

    name = (data or {}).get('name', '').strip()
    email = (data or {}).get('email', '').strip()
    message = (data or {}).get('message', '').strip()

    if not name or not email or not message:
        return jsonify({'error': 'Missing fields'}), 400
    if not is_valid_email(email):
        return jsonify({'error': 'Invalid email address'}), 400
    if len(name) > 100 or len(email) > 100 or len(message) > 2000:
        return jsonify({'error': 'Input too long'}), 400

    subject = f'Contact Form: {name}'
    content = f"Name: {name}\nEmail: {email}\nMessage: {message}"

    mail = Mail(from_email=VERIFIED_SENDER, to_emails=RECEIVER_EMAIL,
                subject=subject, html_content=content)
    mail.reply_to = Email(email)

    try:
        sg = SendGridAPIClient(SENDGRID_API_KEY)
        response = sg.send(mail)
        if response.status_code >= 400:
            logger.error("SendGrid error: %s", response.body)
            return jsonify({'error': 'Failed to send email'}), 500
    except Exception:
        logger.exception("SendGrid send failed")
        return jsonify({'error': 'Failed to send email'}), 500

    return jsonify({'success': True}), 202


# ----------------------------------------------------------------------------------------------------------------------
# File Converter Endpoint
# ----------------------------------------------------------------------------------------------------------------------
@app.route('/api/convert', methods=['POST'])
@limiter.limit('5 per minute') if USE_LIMITER else (lambda f: f)
def convert_file():
    logger.info("POST /api/convert")

    if 'file' not in request.files or 'format' not in request.form:
        return jsonify({'error': 'Missing file or format'}), 400

    target_format = request.form['format'].lower()
    uploaded_file = request.files['file']
    filename = uploaded_file.filename.lower()

    if not filename:
        return jsonify({'error': 'Invalid file'}), 400

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as temp_in:
            uploaded_file.save(temp_in.name)
            input_path = temp_in.name

        output_suffix = f".{target_format}"
        output_path = tempfile.NamedTemporaryFile(delete=False, suffix=output_suffix).name

        # Document → PDF
        if filename.endswith(('.txt', '.md', '.html', '.docx')) and target_format == 'pdf':
            if filename.endswith('.md') or filename.endswith('.txt'):
                with open(input_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                html = markdown(text) if filename.endswith('.md') else f"<pre>{text}</pre>"
            elif filename.endswith('.html'):
                with open(input_path, 'r', encoding='utf-8') as f:
                    html = f.read()
            elif filename.endswith('.docx'):
                doc = Document(input_path)
                html = ''.join(f"<p>{para.text}</p>" for para in doc.paragraphs)

            config = pdfkit.configuration(wkhtmltopdf='/usr/bin/wkhtmltopdf')
            pdfkit.from_string(html, output_path, configuration=config)
            logger.info("Converted document to PDF")

        # DOCX → TXT
        elif filename.endswith('.docx') and target_format == 'txt':
            doc = Document(input_path)
            with open(output_path, 'w', encoding='utf-8') as out:
                for para in doc.paragraphs:
                    out.write(para.text + '\n')
            logger.info("Converted DOCX to TXT")

        # Audio conversion
        elif filename.endswith(('.mp3', '.wav', '.ogg')):
            audio = AudioSegment.from_file(input_path)
            audio.export(output_path, format=target_format)
            logger.info("Converted audio to %s", target_format)

        else:
            return jsonify({'error': 'Unsupported file or conversion type'}), 400

        return send_file(output_path, as_attachment=True)

    except Exception:
        logger.exception("File conversion failed")
        return jsonify({'error': 'Conversion failed'}), 500


# ----------------------------------------------------------------------------------------------------------------------
# JSON Validator Endpoint
# ----------------------------------------------------------------------------------------------------------------------
@app.route('/api/validate-json', methods=['POST'])
def validate_json():
    logger.info("POST /api/validate-json")
    data = request.get_json(force=True, silent=True)
    print("DATA RECEIVED:", data)

    if not data or 'json' not in data:
        return jsonify({'valid': False, 'error': 'Missing "json" field'}), 400

    raw = data['json']

    if not raw.strip():
        return jsonify({'valid': False, 'error': 'Input is empty'}), 200

    try:
        parsed = json.loads(raw)
        formatted = json.dumps(parsed, indent=2)
        return jsonify({'valid': True, 'formatted': formatted}), 200
    except Exception as e:
        return jsonify({'valid': False, 'error': str(e)}), 200


# ----------------------------------------------------------------------------------------------------------------------
# Markdown Previewer Endpoint
# ----------------------------------------------------------------------------------------------------------------------
@app.route('/api/markdown', methods=['POST'])
def render_markdown():
    logger.info("POST /api/markdown")

    data = request.get_json(force=True, silent=True)
    if not data or 'markdown' not in data:
        return jsonify({'error': 'Missing markdown field'}), 400

    md = data['markdown']
    if not isinstance(md, str):
        return jsonify({'error': 'Invalid markdown input'}), 400
    if len(md) > 20000:
        return jsonify({'error': 'Markdown too long (20,000 character limit)'}), 400

    # Allow options/extras, but only from a safe whitelist
    allowed_extras = {
        "tables", "fenced-code-blocks", "strike", "cuddled-lists",
        "break-on-newline", "autolink", "header-ids", "safe-links"
    }
    user_extras = data.get('options', [])
    if not isinstance(user_extras, list):
        user_extras = []
    selected_extras = [e for e in user_extras if e in allowed_extras]
    if not selected_extras:
        selected_extras = ["fenced-code-blocks", "strike", "tables", "cuddled-lists", "break-on-newline", "safe-links"]

    html = markdown(
        md,
        extras=selected_extras,
        safe_mode="escape"
    )

    allowed_tags = set(bleach.sanitizer.ALLOWED_TAGS).union({
        'p', 'pre', 'code', 'hr', 'br', 'span', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table', 'thead', 'tbody', 'tr',
        'th', 'td'
    })
    allowed_attrs = {
        'a': ['href', 'title', 'rel'],
        'img': ['src', 'alt', 'title'],
        'span': ['class'],
        'code': ['class'],
        'th': ['colspan', 'rowspan'],
        'td': ['colspan', 'rowspan'],
    }
    allowed_protocols = set(bleach.sanitizer.ALLOWED_PROTOCOLS).union({'data'})

    clean_html = bleach.clean(
        html,
        tags=allowed_tags,
        attributes=allowed_attrs,
        protocols=allowed_protocols,
        strip=True
    )

    return jsonify({'html': clean_html}), 200


# ----------------------------------------------------------------------------------------------------------------------
# URLShorten Endpoint
# ----------------------------------------------------------------------------------------------------------------------
short_to_url = {}
url_to_short = {}


def is_private_url(url):
    try:
        hostname = urlparse(url).hostname
        ip = socket.gethostbyname(hostname)
        return ipaddress.ip_address(ip).is_private
    except Exception:
        return True  # treat unknowns as unsafe


def generate_code(length=6):
    chars = string.ascii_letters + string.digits
    return ''.join(random.choices(chars, k=length))


@app.route('/api/shorten', methods=['POST'])
def shorten_url():
    logger.info("POST /api/shorten")

    data = request.get_json(force=True, silent=True)
    if not data or 'url' not in data:
        return jsonify({'error': 'Missing URL'}), 400

    url = data['url'].strip()
    if not (url.startswith('http://') or url.startswith('https://')):
        return jsonify({'error': 'URL must start with http:// or https://'}), 400
    if len(url) > 2048:
        return jsonify({'error': 'URL too long'}), 400

    db = get_db()
    # Check for existing URL
    row = db.execute('SELECT code FROM short_urls WHERE url = ?', (url,)).fetchone()
    if row:
        code = row['code']
    else:
        # Generate a new unique code
        while True:
            code = generate_code()
            if not db.execute('SELECT 1 FROM short_urls WHERE code = ?', (code,)).fetchone():
                break
        db.execute('INSERT INTO short_urls (code, url) VALUES (?, ?)', (code, url))
        db.commit()

    return jsonify({'code': code}), 200


@app.route('/s/<code>')
def redirect_short_url(code):
    logger.info(f"Redirect /s/{code}")

    db = get_db()
    row = db.execute('SELECT url FROM short_urls WHERE code = ?', (code,)).fetchone()
    if row:
        return redirect(row['url'])
    else:
        return "<h2>Invalid or expired link.</h2>", 404


# ----------------------------------------------------------------------------------------------------------------------
# Chat Endpoint
# ----------------------------------------------------------------------------------------------------------------------
@app.route('/chat', methods=['POST'])
def chat():
    logger.info("POST /chat")

    data = request.get_json(force=True, silent=True)
    message = (data or {}).get('message', '').strip()
    user_agent = request.headers.get('User-Agent', 'Unknown')
    ip = request.remote_addr

    if not message:
        return jsonify({'error': 'Missing message'}), 400

    if len(message) > 1000 or len(message) < 1:
        return jsonify({'error': 'Invalid message length'}), 400

    content = (
        f"**New Website Chat Message**\n"
        f"Message: {message}\n"
        f"IP: `{ip}`\n"
        f"User-Agent: `{user_agent}`"
    )

    if not DISCORD_WEBHOOK_URL:
        logger.error("DISCORD_WEBHOOK_URL is not set")
        return jsonify({'error': 'Webhook URL not set'}), 500

    try:
        resp = requests.post(
            DISCORD_WEBHOOK_URL,
            json={"content": content}
        )
        if resp.status_code >= 400:
            logger.error(f"Discord response: {resp.status_code} {resp.text}")
            return jsonify({'error': 'Failed to deliver to Discord', 'details': resp.text}), 502
    except Exception as e:
        logger.exception("Failed to send to Discord webhook")
        return jsonify({'error': 'Internal error', 'details': str(e)}), 500

    return jsonify({'success': True})


# ----------------------------------------------------------------------------------------------------------------------
# Stripe & Email Endpoints (Cleaned Up)
# ----------------------------------------------------------------------------------------------------------------------
stripe.api_key = STRIPE_SECRET_KEY


def send_email_with_download_link(email, product_id):
    download_url = f"{FRONTEND_BASE_URL}/download/{product_id}"
    subject = "Your Module Download"
    content = (
        f"Thank you for your purchase!<br><br>"
        f"You can download your file using the link below:<br>"
        f"<a href='{download_url}'>Download your module</a><br><br>"
        f"If you have any issues, reply to this email."
    )
    mail = Mail(
        from_email=VERIFIED_SENDER,
        to_emails=email,
        subject=subject,
        html_content=content
    )
    sg = SendGridAPIClient(SENDGRID_API_KEY)
    sg.send(mail)


def send_game_server_instructions(email, server_id):
    subject = "Game Server Purchase Instructions"
    content = (
        f"Thank you for your purchase of {server_id} server hosting!<br><br>"
        f"Please reply to this email or join our Discord to complete your setup and get access to your server.<br>"
        f"Instructions:<br>"
        f"1. Join: <a href='https://discord.gg/u6GqqWxqRN'>Discord</a><br>"
        f"2. Open a ticket with your order info.<br>"
        f"3. We'll get you set up within a few hours.<br>"
    )
    mail = Mail(
        from_email=VERIFIED_SENDER,
        to_emails=email,
        subject=subject,
        html_content=content
    )
    sg = SendGridAPIClient(SENDGRID_API_KEY)
    sg.send(mail)


@app.route('/api/create-checkout-session', methods=['POST'])
def create_checkout_session():
    logger.info("POST /api/create-checkout-session")

    data = request.get_json(force=True, silent=True)
    product_type = data.get('type')  # "module" or "gameServer"
    label = data.get('label')
    price = data.get('price')
    description = data.get('desc', '')
    product_id = data.get('productId')  # file to deliver (for modules)
    server_id = data.get('productId')  # game server identifier (for gameServers)
    customer_email = data.get('email')  # optionally collected

    if not label or not price or not product_type:
        return jsonify({'error': 'Missing required fields'}), 400

    currency = 'gbp'

    try:
        if product_type == "gameServer":
            amount = int(float(price.replace('£', '').replace('/mo', '').strip()))
            unit_amount = amount * 100
            recurring = {"interval": "month"}
        else:
            amount = int(float(price.replace('£', '').replace('+', '').strip()))
            unit_amount = amount * 100
            recurring = None
    except Exception:
        return jsonify({'error': 'Invalid price format'}), 400

    try:
        if product_type == "gameServer":
            product = stripe.Product.create(name=label, description=description)
            price_obj = stripe.Price.create(
                unit_amount=unit_amount,
                currency=currency,
                recurring=recurring,
                product=product.id
            )
            mode = 'subscription'
            line_items = [{'price': price_obj.id, 'quantity': 1}]
            metadata = {
                'product_type': product_type,
                'server_id': server_id
            }
        else:
            line_items = [{
                'price_data': {
                    'currency': currency,
                    'product_data': {
                        'name': label,
                        'description': description,
                    },
                    'unit_amount': unit_amount,
                },
                'quantity': 1,
            }]
            mode = 'payment'
            metadata = {
                'product_type': product_type,
                'product_id': product_id or label
            }

        session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=line_items,
            mode=mode,
            success_url=f'{FRONTEND_BASE_URL}/thank-you?session_id={{CHECKOUT_SESSION_ID}}',
            cancel_url=f'{FRONTEND_BASE_URL}/cancelled',
            customer_email=customer_email if customer_email else None,
            metadata=metadata
        )
        return jsonify({'url': session.url})
    except Exception as e:
        logger.exception("Stripe checkout session creation failed")
        return jsonify({'error': str(e)}), 500


@app.route('/api/stripe/webhook', methods=['POST'])
def stripe_webhook():
    logger.info("POST /api/stripe/webhook")

    payload = request.data
    sig_header = request.headers.get('stripe-signature')
    try:
        event = stripe.Webhook.construct_event(
            payload, sig_header, STRIPE_WEBHOOK_SECRET
        )
    except (ValueError, stripe.error.SignatureVerificationError):
        return '', 400

    if event['type'] in ['checkout.session.completed', 'checkout.session.async_payment_succeeded']:
        session = event['data']['object']
        customer_email = (
                session.get('customer_email') or
                (session.get('customer_details', {}).get('email')) or
                None
        )
        metadata = session.get('metadata', {})
        product_type = metadata.get('product_type')
        if product_type == 'module':
            product_id = metadata.get('product_id')
            send_email_with_download_link(customer_email, product_id)
        elif product_type == 'gameServer':
            server_id = metadata.get('server_id')
            send_game_server_instructions(customer_email, server_id)

    return '', 200


@app.route('/download/<product_id>')
def download_product(product_id):
    session_id = request.args.get('session_id')
    if not session_id:
        return abort(403, "Missing session id")

    # 1. Fetch the session from Stripe
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        created = session.get('created', 0)
        metadata = session.get('metadata', {})
        allowed_product = metadata.get('product_id')
        product_type = metadata.get('product_type')
    except Exception:
        return abort(403, "Invalid session id")

    # 2. Must be module type and correct file
    if product_type != "module" or allowed_product != product_id:
        return abort(403, "Not authorized for this download")

    # 3. Only allow if <1 hour old
    if time.time() - created > 3600:
        return abort(403, "This download link has expired. Please open a support ticket.")

    allowed_files = set(os.listdir('products'))
    if product_id not in allowed_files:
        return "Not found", 404
    return send_file(os.path.join('products', product_id), as_attachment=True)


@app.route('/api/stripe/session-info')
def stripe_session_info():
    logger.info("GET /api/stripe/session-info")

    session_id = request.args.get('session_id')
    if not session_id:
        return jsonify({'error': 'Missing session_id'}), 400
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        metadata = session.get('metadata', {})
        label = metadata.get('label') or ""
        customer_email = session.get('customer_email')
        if not customer_email and session.get('customer'):
            customer_obj = stripe.Customer.retrieve(session.get('customer'))
            customer_email = customer_obj.get('email')
        created = session.get('created', 0)  # UNIX timestamp
        return jsonify({
            'product_type': metadata.get('product_type'),
            'product_id': metadata.get('product_id'),
            'label': label,
            'customer_email': customer_email,
            'created': created,
        })
    except Exception as e:
        logger.exception("Stripe session info fetch failed")
        return jsonify({'error': 'Could not fetch session info.'}), 400


# ----------------------------------------------------------------------------------------------------------------------
# Yugioh Endpoints
# ----------------------------------------------------------------------------------------------------------------------


def get_ebay_prices(card_name, region='GB'):
    url = "https://api.ebay.com/buy/browse/v1/item_summary/search"
    marketplace_id = 'EBAY_GB' if region == 'GB' else 'EBAY_US'
    params = {
        'q': card_name,
        'limit': 5,
        'sort': 'price',
        'filter': f"itemLocationCountry:{region}"
    }

    access_token = get_ebay_access_token()
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json',
        'X-EBAY-C-MARKETPLACE-ID': marketplace_id,
    }
    response = requests.get(url, headers=headers, params=params)
    try:
        data = response.json()
    except Exception:
        return []
    results = []
    for item in data.get('itemSummaries', []):
        price = item.get('price', {})
        price_str = f"{price.get('value')} {price.get('currency')}" if price else "N/A"
        results.append({
            'title': item.get('title'),
            'price': price_str,
            'condition': item.get('condition', 'Unknown'),
            'url': item.get('itemWebUrl', '')
        })
    return results


def get_ebay_access_token():
    global ebay_access_token, ebay_token_expiry
    now = time.time()
    # If we have a token and it's not expiring in the next 2 minutes, use it
    if ebay_access_token and now < ebay_token_expiry - 120:
        return ebay_access_token

    # Else, refresh the token
    import requests, base64
    credentials = f"{EBAY_CLIENT_ID}:{EBAY_CLIENT_SECRET}"
    b64credentials = base64.b64encode(credentials.encode()).decode()
    headers = {
        "Content-Type": "application/x-www-form-urlencoded",
        "Authorization": f"Basic {b64credentials}"
    }
    data = {
        "grant_type": "refresh_token",
        "refresh_token": EBAY_REFRESH_TOKEN,
        "scope": "https://api.ebay.com/oauth/api_scope"
    }
    resp = requests.post("https://api.ebay.com/identity/v1/oauth2/token", headers=headers, data=data)
    if resp.ok:
        tokens = resp.json()
        ebay_access_token = tokens["access_token"]
        ebay_token_expiry = now + int(tokens.get("expires_in", 7200))
        return ebay_access_token
    else:
        raise Exception(f"Failed to refresh eBay access token: {resp.text}")

@app.route('/api/exchange-rates')
def exchange_rates():
    # Hardcoded for demo; ideally fetch from a live API and cache hourly
    rates = {
        "USD_GBP": 0.79,
        "EUR_GBP": 0.85,
        "USD_USD": 1.0,
        "EUR_USD": 1.09,
    }
    return jsonify(rates)

@app.route('/api/yugioh/cards', methods=['POST'])
def get_card_data():
    logger.info("POST /api/yugioh/cards")
    data = request.get_json(force=True, silent=True)
    ids = data.get('ids', [])
    if not isinstance(ids, list) or not all(isinstance(i, str) for i in ids):
        return jsonify({'error': 'Invalid or missing ID list'}), 400
    results = [CARD_LOOKUP[i] for i in ids if i in CARD_LOOKUP]
    return jsonify({'data': results})


@app.route('/api/yugioh/search', methods=['GET'])
def search_cards():
    logger.info(f"GET /api/yugioh/search q={request.args.get('q', '').strip()}")
    query = request.args.get('q', '').strip().lower()
    if not query or len(query) < 2:
        return jsonify({'results': []})
    matches = [card for card in ALL_CARDS if query in card.get('name', '').lower()]
    return jsonify({'results': matches[:25]})


@app.route('/api/yugioh/ebay-xlsx/start', methods=['POST'])
def ebay_xlsx_start():
    logger.info("POST /api/yugioh/ebay-xlsx/start")
    req = request.get_json(force=True, silent=True)
    deck = req.get('deck', {})
    currency = req.get('currency', 'GBP').upper()
    job_id = str(uuid.uuid4())
    generation_jobs[job_id] = {'progress': 0, 'ready': False, 'file': None, 'error': None}
    total = sum(len(deck.get(section, [])) for section in ['main', 'extra', 'side'])
    if total == 0:
        generation_jobs[job_id]['ready'] = True
        generation_jobs[job_id]['error'] = "Deck is empty."
        return jsonify({'job_id': job_id}), 202

    def group_cards(cards):
        grouped = {}
        for card in cards:
            key = card['name']
            grouped.setdefault(key, {'name': card['name'], 'count': 0, **card})
            grouped[key]['count'] += 1
        return list(grouped.values())

    def do_work():
        logging.info(
            f"Starting XLSX generation for job_id={job_id}, region={currency}, deck size: {sum(len(deck.get(section, [])) for section in ['main', 'extra', 'side'])}")
        try:
            region = 'US' if currency == 'USD' else 'GB'
            output = io.BytesIO()
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "eBay Prices"

            section_fill = PatternFill(start_color="A6CAF0", end_color="A6CAF0", fill_type="solid")
            header_fill = PatternFill(start_color="B7D3F2", end_color="B7D3F2", fill_type="solid")
            border = Border(left=Side(style='thin'), right=Side(style='thin'),
                            top=Side(style='thin'), bottom=Side(style='thin'))
            bold = Font(bold=True)
            big_bold = Font(bold=True, size=13)
            center = Alignment(horizontal='center', vertical='center')
            left = Alignment(horizontal='left', vertical='center')

            columns = [("Card", 25), ("Set Codes", 20), ("Quantity", 10), ("Price", 13), ("URL", 45), ("Total", 12)]
            section_order = [('main', 'Main Deck'), ('extra', 'Extra Deck'), ('side', 'Side Deck')]
            currency_format = '"$"#,##0.00' if currency == "USD" else '"£"#,##0.00'
            col_base = 2
            row = 2
            section_total_cells = []

            # --- Collect all unique card queries
            all_card_queries = {}
            for section, _ in section_order:
                for card in group_cards(deck.get(section, [])):
                    setcodes = []
                    if 'card_sets' in card and card['card_sets']:
                        setcodes = [s['set_code'] for s in card['card_sets'] if 'EN' in s['set_code']]
                    all_card_queries[card['name']] = setcodes

            # --- Prepare price lookups in parallel
            price_results = {}

            def lookup_card(card_name, setcodes):
                best_price = None
                best_url = ""
                setcodes_str = ", ".join(setcodes) if setcodes else ""
                # Query eBay for each EN setcode, use lowest price
                for code in setcodes:
                    query = f"{code} {card_name}"
                    ebay_results = get_ebay_prices(query, region=region)
                    filtered = [x for x in ebay_results if x['price'] != "N/A"]
                    if filtered:
                        def price_to_float(p):
                            try:
                                return float(p['price'].replace('£', '').replace('$', '').split()[0])
                            except Exception:
                                return float('inf')

                        cheapest = min(filtered, key=price_to_float)
                        price = price_to_float(cheapest)
                        if best_price is None or price < best_price:
                            best_price = price
                            best_url = cheapest['url']
                # If no EN setcode found, fallback to card name
                if best_price is None:
                    ebay_results = get_ebay_prices(card_name, region=region)
                    filtered = [x for x in ebay_results if x['price'] != "N/A"]
                    if filtered:
                        def price_to_float(p):
                            try:
                                return float(p['price'].replace('£', '').replace('$', '').split()[0])
                            except Exception:
                                return float('inf')

                        cheapest = min(filtered, key=price_to_float)
                        best_price = price_to_float(cheapest)
                        best_url = cheapest['url']
                return (setcodes_str, best_price if best_price is not None else 0.0, best_url)

            with ThreadPoolExecutor(max_workers=8) as executor:
                future_to_name = {
                    executor.submit(lookup_card, card_name, setcodes): card_name
                    for card_name, setcodes in all_card_queries.items()
                }
                num_cards = len(all_card_queries)
                extra_steps = len(section_order) + 2
                progress = 0
                progress_steps = num_cards + extra_steps
                max_percent = 0

                for future in as_completed(future_to_name):
                    card_name = future_to_name[future]
                    try:
                        price_results[card_name] = future.result()
                    except Exception as e:
                        logging.exception(f"Failed to fetch price for {card_name}")
                        price_results[card_name] = ("", 0.0, "")
                    # progress per card
                    progress += 1
                    percent = int(100 * progress / progress_steps)
                    max_percent = max(percent, max_percent)
                    generation_jobs[job_id]['progress'] = max_percent

            # --- Now, generate the sheet as before, using the fetched prices
            for section, label in section_order:
                cards = group_cards(deck.get(section, []))
                ws.merge_cells(start_row=row, start_column=col_base, end_row=row,
                               end_column=col_base + len(columns) - 1)
                cell = ws.cell(row=row, column=col_base, value=label)
                cell.fill = section_fill
                cell.font = big_bold
                cell.alignment = center
                row += 1

                # Headers
                for idx, (head, _) in enumerate(columns, start=col_base):
                    c = ws.cell(row=row, column=idx, value=head)
                    c.fill = header_fill
                    c.font = bold
                    c.alignment = center
                    c.border = border

                row += 1
                card_start_row = row

                for card in cards:
                    setcodes_str, best_price, best_url = price_results.get(card['name'], ("", 0.0, ""))

                    ws.cell(row=row, column=col_base, value=card['name']).alignment = left
                    ws.cell(row=row, column=col_base).border = border

                    ws.cell(row=row, column=col_base + 1, value=setcodes_str).alignment = left
                    ws.cell(row=row, column=col_base + 1).border = border

                    ws.cell(row=row, column=col_base + 2, value=card['count']).alignment = center
                    ws.cell(row=row, column=col_base + 2).border = border

                    price_cell = ws.cell(row=row, column=col_base + 3,
                                         value=best_price if best_price is not None else 0.0)
                    price_cell.number_format = currency_format
                    price_cell.alignment = center
                    price_cell.border = border

                    ws.cell(row=row, column=col_base + 4, value=best_url).alignment = left
                    ws.cell(row=row, column=col_base + 4).border = border

                    quantity_ref = f"{get_column_letter(col_base + 2)}{row}"
                    price_ref = f"{get_column_letter(col_base + 3)}{row}"
                    total_cell = ws.cell(row=row, column=col_base + 5)
                    total_cell.value = f"=SUM({quantity_ref}*{price_ref})"
                    total_cell.number_format = currency_format
                    total_cell.alignment = center
                    total_cell.border = border

                    row += 1

                # Section totals
                ws.cell(row=row, column=col_base + 4, value="Total").font = bold
                ws.cell(row=row, column=col_base + 4).alignment = center
                ws.cell(row=row, column=col_base + 4).border = border
                sum_formula = f"=SUM({get_column_letter(col_base + 5)}{card_start_row}:{get_column_letter(col_base + 5)}{row - 1})" if cards else ""
                total_sum_cell = ws.cell(row=row, column=col_base + 5, value=sum_formula)
                total_sum_cell.font = bold
                total_sum_cell.number_format = currency_format
                total_sum_cell.alignment = center
                total_sum_cell.border = border
                section_total_cells.append(total_sum_cell.coordinate)
                row += 2

                inc = random.randint(4, 7)
                progress += 1
                percent = min(int(100 * progress / progress_steps) + inc, 99)
                max_percent = max(percent, max_percent)
                generation_jobs[job_id]['progress'] = max_percent

            # Grand total row (randomize fake step)
            ws.cell(row=row, column=col_base + 4, value="Total").font = big_bold
            ws.cell(row=row, column=col_base + 4).alignment = center
            ws.cell(row=row, column=col_base + 4).border = border
            if section_total_cells:
                sum_formula = f"=SUM({','.join(section_total_cells)})"
                total_cell = ws.cell(row=row, column=col_base + 5, value=sum_formula)
                total_cell.font = big_bold
                total_cell.number_format = currency_format
                total_cell.alignment = center
                total_cell.border = border

            inc = random.randint(3, 5)
            progress += 1
            percent = min(int(100 * progress / progress_steps) + inc, 99)
            max_percent = max(percent, max_percent)
            generation_jobs[job_id]['progress'] = max_percent

            for idx, (_, width) in enumerate(columns, start=col_base):
                ws.column_dimensions[get_column_letter(idx)].width = width

            wb.save(output)
            output.seek(0)
            inc = random.randint(1, 4)
            progress += 1
            percent = min(int(100 * progress / progress_steps) + inc, 99)
            max_percent = max(percent, max_percent)
            generation_jobs[job_id]['progress'] = max_percent

            generation_jobs[job_id]['file'] = output.read()
            generation_jobs[job_id]['ready'] = True
            generation_jobs[job_id]['progress'] = 100
        except Exception as e:
            generation_jobs[job_id]['error'] = str(e)
            generation_jobs[job_id]['ready'] = True
            generation_jobs[job_id]['progress'] = 100

    threading.Thread(target=do_work, daemon=True).start()
    return jsonify({'job_id': job_id}), 202


@app.route('/api/yugioh/ebay-xlsx/progress')
def ebay_xlsx_progress():
    job_id = request.args.get('job_id')
    job = generation_jobs.get(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404
    return jsonify({
        'progress': job['progress'],
        'ready': job['ready'],
        'error': job['error']
    })


@app.route('/api/yugioh/ebay-xlsx/download')
def ebay_xlsx_download():
    logger.info(f"GET /api/yugioh/ebay-xlsx/download job_id={request.args.get('job_id')}")
    job_id = request.args.get('job_id')
    job = generation_jobs.get(job_id)
    if not job or not job.get('ready') or not job.get('file'):
        return jsonify({'error': 'Not ready'}), 400
    return send_file(
        io.BytesIO(job['file']),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='ebay_prices.xlsx'
    )


@app.route('/api/yugioh/ebay-prices', methods=['GET'])
def yugioh_ebay_prices():
    logger.info(f"GET /api/yugioh/ebay-prices name={request.args.get('name', '').strip()} region={request.args.get('region', 'GB').upper()}")
    card_name = request.args.get('name', '').strip()
    region = request.args.get('region', 'GB').upper()
    if not card_name:
        return jsonify({'error': 'Missing card name'}), 400
    results = get_ebay_prices(card_name, region=region)
    return jsonify({'results': results})


# ----------------------------------------------------------------------------------------------------------------------
# Start Server
# ----------------------------------------------------------------------------------------------------------------------
if USE_LIMITER:
    limiter.exempt(ebay_xlsx_progress)
    limiter.exempt(render_markdown)
    limiter.exempt(search_cards)

# Start background card update loop
threading.Thread(target=update_cards_periodically, daemon=True).start()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5010, debug=True)
